import os
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from docx import Document
from bs4 import BeautifulSoup
import pandas as pd

app = Flask(__name__)
app.secret_key = 'secret_key'  # Necessário para usar flash messages

# Define o diretório de upload
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Página inicial do gerenciador
@app.route('/')
def index():
    arquivos = os.listdir(app.config['UPLOAD_FOLDER'])
    return render_template('index.html', arquivos=arquivos, get_file_icon=get_file_icon)

# Rota para upload de arquivos
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('Nenhum arquivo selecionado.')
        return redirect(request.url)
    file = request.files['file']
    if file.filename != '':
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))
        flash(f'O arquivo {file.filename} foi carregado com sucesso!')
    return redirect(url_for('index'))

# Rota para excluir arquivos
@app.route('/delete/<filename>', methods=['POST'])
def delete_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        flash(f'O arquivo {filename} foi excluído com sucesso!')
    else:
        flash('Arquivo não encontrado.')
    return redirect(url_for('index'))

# Rota para download de arquivos
@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

# Função para identificar o tipo de arquivo
def get_file_type(filename):
    ext = filename.split('.')[-1].lower()
    if ext in ['txt', 'md', 'csv', 'html', 'css', 'js', 'py']:
        return 'text'
    elif ext == 'docx':
        return 'docx'
    elif ext == 'xlsx':
        return 'spreadsheet'
    elif ext in ['png', 'jpg', 'jpeg', 'gif']:
        return 'image'
    else:
        return 'unsupported'

def get_file_icon(filename):
    ext = filename.split('.')[-1].lower()
    icons = {
        'txt': 'text-icon.png',
        'md': 'text-icon.png',
        'csv': 'spreadsheet-icon.png',
        'html': 'html-icon.png',
        'css': 'css-icon.png',
        'js': 'js-icon.png',
        'py': 'python-icon.png',
        'docx': 'word-icon.png',
        'xlsx': 'excel-icon.png',
        'png': 'image-icon.png',
        'jpg': 'image-icon.png',
        'jpeg': 'image-icon.png',
        'gif': 'image-icon.png'
    }
    return icons.get(ext, 'unknown-icon.png')

# Rota para abrir um arquivo para edição ou visualização
@app.route('/edit/<filename>', methods=['GET'])
def edit_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_type = get_file_type(filename)
    
    if file_type == 'text':
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return render_template('edit_text.html', filename=filename, content=content)
    elif file_type == 'docx':
        if os.path.exists(file_path):
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            return render_template('edit_text.html', filename=filename, content=content)
    elif file_type == 'spreadsheet':
        if os.path.exists(file_path):
            df = pd.read_excel(file_path)
            content = df.to_html()
            return render_template('view_spreadsheet.html', filename=filename, content=content)
    elif file_type == 'image':
        return render_template('view_image.html', filename=filename)
    else:
        flash('Tipo de arquivo não suportado para edição.')
        return redirect(url_for('index'))

# Rota para salvar o arquivo editado
@app.route('/save/<filename>', methods=['POST'])
def save_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_type = get_file_type(filename)
    
    if file_type == 'text':
        new_content = request.form['content']
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(new_content)
        flash(f'O arquivo {filename} foi salvo com sucesso!')
    
    elif file_type == 'docx':
        # Extrair o texto do conteúdo HTML usando BeautifulSoup para remover tags HTML
        new_content = request.form['content']
        soup = BeautifulSoup(new_content, "html.parser")
        plain_text = soup.get_text()

        doc = Document()
        for line in plain_text.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
        flash(f'O arquivo {filename} foi salvo com sucesso!')
    
    elif file_type == 'spreadsheet':
        content = request.get_json().get('content')
        df = pd.DataFrame(content)
        df.to_excel(file_path, index=False)
        flash(f'O arquivo {filename} foi salvo com sucesso!')
    
    return redirect(url_for('index'))

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)
